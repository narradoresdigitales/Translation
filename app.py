import streamlit as st
from docx import Document
from deep_translator import GoogleTranslator
from io import BytesIO
from datetime import datetime

# -----------------------------
# Page Config
# -----------------------------
st.set_page_config(
    page_title="Document Translator",
    layout="wide"
)

st.title("📄 Word Document Translator")

# -----------------------------
# Language Options
# -----------------------------
LANGUAGES = {
    "English": "english",
    "Spanish": "spanish",
    "French": "french",
    "German": "german",
    "Italian": "italian",
    "Portuguese": "portuguese",
    "Russian": "russian",
    "Arabic": "arabic",
    "Korean": "korean",
    "Chinese (Simplified)": "chinese (simplified)",
    "Chinese (Traditional)": "chinese (traditional)",
}

# -----------------------------
# Helper Functions
# -----------------------------
def extract_docx_paragraphs(file):
    doc = Document(file)
    return [p.text for p in doc.paragraphs if p.text.strip()]


def translate_paragraphs(paragraphs, target_language):
    translator = GoogleTranslator(source="auto", target=target_language)
    translated = []
    for p in paragraphs:
        translated.append(translator.translate(p))
    return translated


def build_translated_docx(paragraphs):
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    return doc


# -----------------------------
# UI Controls
# -----------------------------
with st.sidebar:
    st.header("Settings")

    language_label = st.selectbox(
        "Translate to:",
        list(LANGUAGES.keys())
    )

    target_language = LANGUAGES[language_label]

    preview_count = st.slider(
        "Preview paragraphs",
        min_value=1,
        max_value=20,
        value=10
    )

uploaded_file = st.file_uploader(
    "Upload a Word document (.docx)",
    type=["docx"]
)

translate_clicked = st.button("🌍 Translate Document")

# -----------------------------
# Main Logic
# -----------------------------
if translate_clicked:
    if not uploaded_file:
        st.error("Please upload a .docx file.")
        st.stop()

    with st.spinner("Extracting document text..."):
        original_paragraphs = extract_docx_paragraphs(uploaded_file)

    if not original_paragraphs:
        st.error("The document contains no readable text.")
        st.stop()

    st.success(f"Extracted {len(original_paragraphs)} paragraphs.")

    with st.spinner("Translating document..."):
        translated_paragraphs = translate_paragraphs(
            original_paragraphs,
            target_language
        )

    # -----------------------------
    # Preview Section
    # -----------------------------
    st.subheader("🔍 Translation Preview")

    col1, col2 = st.columns(2)

    with col1:
        st.markdown("### Original")
        for p in original_paragraphs[:preview_count]:
            st.markdown(f"> {p}")

    with col2:
        st.markdown("### Translated")
        for p in translated_paragraphs[:preview_count]:
            st.markdown(f"> {p}")

    # -----------------------------
    # Build Downloadable Document
    # -----------------------------
    with st.spinner("Building translated document..."):
        translated_doc = build_translated_docx(translated_paragraphs)
        buffer = BytesIO()
        translated_doc.save(buffer)
        buffer.seek(0)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_filename = f"translated_{target_language}_{timestamp}.docx"

    st.download_button(
        label="⬇️ Download translated document",
        data=buffer,
        file_name=output_filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
